from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("构建更公平的教育未来_课堂展示逐字稿.docx")
FONT_NAME = "Hiragino Sans GB"


TITLE = "构建更公平的教育未来"
SUBTITLE = "全球教育公平数据可视化分析课堂展示逐字稿"


SECTIONS = [
    (
        "开场与项目引入",
        [
            "各位老师、同学大家好。今天我展示的项目是《构建更公平的教育未来：全球教育公平数据可视化分析》。",
            "这个项目关注的是一个看似简单、但其实非常复杂的问题：全球教育公平到底公平在哪里，又不公平在哪里？我们通常会说，教育是改变命运的重要工具，但如果从全球视角来看，不同国家在教育投入、入学机会、师资质量和学习成果上，其实存在非常明显的差距。",
            "所以本项目没有只看单一指标，而是把教育公平拆成三个层面：第一是“投入与师资”，也就是教育资源有没有投入进去，以及这些投入有没有转化为教师质量；第二是“机会与平等”，关注儿童是否真的有机会进入学校，以及性别之间是否平等；第三是“质量与成果”，进一步追问，入学是否就等于真正学到了东西。",
            "现在大家看到的是网页首页。整体设计采用蓝白色调，蓝色代表教育、理性和公共议题，白色背景让图表更加清晰。左侧是目录导航，可以快速切换到三个主题模块。在课堂展示时，我可以把侧边目录折叠，让主体图表居中显示，这样更适合大屏幕展示。",
        ],
        ["展示首页；如需全屏讲解，可先折叠左侧目录。"],
    ),
    (
        "第一部分：投入与师资",
        [
            "接下来我进入第一部分，“投入与师资”。这一部分想回答的问题是：教育投入能否有效转化为师资质量？",
            "首先我们看全球教育经费投入分布图。这里的指标是教育支出占 GDP 的比例。地图颜色越深，代表该国教育支出占 GDP 的比例越高。地图适合展示空间分布，我们可以很直观地看到，不同地区的教育投入并不均衡。",
            "右侧排行榜则补充展示极值国家。通过切换“投入最高”和“投入最低”，可以快速看到哪些国家教育支出比例较高，哪些国家较低。排行榜和地图结合起来，一个负责看整体空间格局，一个负责定位具体国家，二者是互补的。",
            "但是，仅仅知道“投入多少”是不够的。教育投入高，并不一定意味着教育质量就高。所以页面继续用散点图展示教育支出占 GDP 比例和受训教师比例之间的关系。",
            "这张散点图中，横轴是教育支出占 GDP 的比例，纵轴是受训教师比例。每一个点代表一个国家。我们可以切换学前、小学和中学三个教育阶段，观察不同阶段的师资情况。",
            "图中可以看到，教育投入和受训教师比例之间并不是简单的线性关系。有些国家投入并不算特别高，但受训教师比例较高；也有一些国家投入较高，但师资培训比例并没有同步提升。这说明教育公平并不只是“花了多少钱”的问题，还涉及资金使用效率、教师培训体系和资源配置方式。",
            "接下来这张图进一步比较不同教育阶段的师资配置。这里用点和连线展示学前、小学和中学三个阶段的受训教师比例。点的位置越靠右，说明受训教师比例越高；连线越长，说明不同教育阶段之间差距越大。",
            "这张图的一个重要发现是：学前教育师资在很多国家是明显短板。相比小学和中学，学前阶段受训教师比例更加不稳定，部分国家差距很大。这提醒我们，教育公平不能只关注义务教育阶段，早期教育同样需要被重视。",
            "所以第一部分的核心结论是：教育投入重要，但投入本身并不自动转化为高质量师资；同时，学前教育师资短板是全球教育资源配置中的一个突出问题。",
        ],
        [
            "进入“投入与师资”。",
            "切换排行榜中的“投入最高”和“投入最低”。",
            "在散点图中切换“学前 / 小学 / 中学”。",
            "在排行榜条形或散点上悬停，展示交互强调效果。",
        ],
    ),
    (
        "第二部分：机会与平等",
        [
            "第二部分是“机会与平等”。这一部分关注的是：儿童是否真的有机会进入学校？男孩和女孩的教育机会是否一致？",
            "首先我们看小学适龄儿童失学率地图。颜色越深，表示失学率越高，也就是仍未进入学校的儿童比例越高。灰色区域表示缺少数据。这个图让我们直观地看到，全球基础教育虽然已经普及了很多，但部分地区仍然存在较高失学率。",
            "右侧排行榜可以切换“较高”和“较低”，用于查看失学率最高或最低的国家。这里也加入了鼠标悬停强调效果，展示时可以把鼠标移到某个条形上，让大家更清楚地看到对应国家和数值。",
            "接下来是按性别比较失学率。这张图采用左右对称的条形图，也可以理解为“蝴蝶图”。左侧表示女孩失学率，右侧表示男孩失学率。这样设计的好处是，两种性别可以在同一水平线上直接比较，差距一眼就能看出来。",
            "页面还支持选择指定国家进行比较。也就是说，我们不只能随机看一批国家，还可以根据课堂讨论需要，输入或选择某些国家，观察它们的男孩和女孩失学率差异。比如在一些国家，女孩失学率明显高于男孩，这可能与社会文化、家庭经济条件、安全环境等因素有关；而在另一些国家，也可能出现男孩失学率更高的情况。因此，性别不平等并不是单一方向的问题，而是具有地区差异和社会背景差异。",
            "然后我们继续看性别平价指数，也就是 GPI。GPI 接近 1，说明男女基本平等；小于 1，通常说明女性处于相对劣势；大于 1，则说明女性相对占优。网页通过散点图比较小学和高等教育阶段的 GPI，可以看到：很多国家在小学阶段已经接近性别平等，但到了高等教育阶段，差异会重新扩大，并呈现更加复杂的格局。",
            "这说明教育机会公平不是“入学一刻”就结束了。即使基础教育入口趋于平等，继续升学和高等教育阶段仍然可能存在新的不平等。",
            "这一部分最后还提供了点击国家查看完成率的交互地图。点击地图上的国家，可以看到该国小学和高等教育完成情况的变化。这里要注意，有些完成率可能超过 100%，这并不一定是错误，而可能与超龄入学、留级、统计口径或人口基数有关。",
            "第二部分的核心结论是：全球基础教育机会整体改善，但部分地区失学率仍然很高；小学阶段性别差距在缩小，但高等教育阶段的性别差异仍然复杂存在。",
        ],
        [
            "进入“机会与平等”。",
            "切换地图查看对象：总体、男孩、女孩。",
            "在失学率排行榜上悬停，强调单个条形。",
            "在性别对比图中选择指定国家，展示国家比较功能。",
            "点击完成率地图上的国家，展示详情面板。",
        ],
    ),
    (
        "第三部分：质量与成果",
        [
            "第三部分是“质量与成果”。这一部分要回答一个更深的问题：进入学校，是否就意味着真正学会了？",
            "这里用两个指标进行分析。第一个是学习贫困率，表示儿童无法达到最低阅读能力标准的比例。第二个是学习调整学年，也就是 LAYS，它不仅考虑学生上了多少年学，还考虑这些学习是否真正转化为知识和能力。",
            "这张散点图中，横轴是学习贫困率，纵轴是学习调整学年。一般来说，学习贫困率越高，LAYS 越低；学习贫困率越低，LAYS 越高。图中的趋势线也体现了这种负相关关系。",
            "这里的交互设计支持搜索和高亮国家。比如我输入一个国家名称，图中对应点会被突出显示，旁边的信息卡会显示这个国家的学习贫困率和 LAYS。颜色分组也可以切换，例如按地区分组或按收入分组，这样可以从不同角度观察国家之间的差异。",
            "右侧是前 20 名 LAYS 排名，用条形图展示学习调整学年较高的国家。条形图适合做排名比较，可以快速看到哪些国家在真实学习成果上表现更好。排序也可以切换，比如按 LAYS 排序或按学习贫困率排序，从而观察不同指标下国家排名的变化。",
            "这一部分最重要的结论是：入学不等于学习。一个国家即使提高了入学率，如果学生没有获得基本阅读能力，没有形成有效学习成果，那么教育公平仍然是不完整的。真正的公平，不仅是让儿童走进学校，还要让他们在学校里获得有质量的学习。",
        ],
        [
            "进入“质量与成果”。",
            "在“高亮国家”中搜索并选择一个国家。",
            "切换颜色分组和条形图排序。",
            "指向右侧前 20 名 LAYS 排名，说明条形图的比较功能。",
        ],
    ),
    (
        "总结收束",
        [
            "最后总结一下。",
            "本项目通过三个层次构建了全球教育公平的分析框架。第一，资源投入层面告诉我们，教育经费很重要，但投入与师资质量之间存在“脱钩”现象，资金使用效率同样关键。第二，机会平等层面告诉我们，全球基础教育普及取得了进展，但失学率和性别差异仍然存在。第三，质量成果层面进一步提醒我们，教育公平不能停留在“是否入学”，还必须关注“是否真正学会”。",
            "从可视化设计上看，本项目使用地图展示空间差异，用排行榜突出极值国家，用散点图揭示变量关系，用对称条形图比较性别差异，并通过搜索、切换、悬停和点击等交互方式，让用户可以从宏观整体进入到具体国家。",
            "所以，这个网页想表达的核心观点是：教育公平不是一个单一指标能够解释的问题。它同时涉及资源、机会和质量。只有当教育投入能够有效转化为师资能力，儿童能够公平进入学校，并且真正获得高质量学习时，我们才更接近一个公平的教育未来。",
            "我的展示到这里结束，谢谢大家。",
        ],
        ["回到首页或停留在总结区，完成展示。"],
    ),
]


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    run.font.size = Pt(10.5)
    run.bold = bold
    run.font.color.rgb = RGBColor(11, 37, 69)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_paragraph_font(paragraph, size=11, color="0B2545", bold=False, italic=False):
    for run in paragraph.runs:
        run.font.name = FONT_NAME
        run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
        run.font.size = Pt(size)
        run.font.color.rgb = RGBColor.from_string(color)
        run.bold = bold
        run.italic = italic


def add_para(doc, text, style=None, size=11, color="0B2545", bold=False, italic=False, after=8):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.333
    run = p.add_run(text)
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic
    return p


def add_instruction(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.right_indent = Inches(0.18)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(f"展示提示：{text}")
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(67, 83, 107)
    run.italic = True
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F4F8FC")
    p_pr.append(shd)


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = FONT_NAME
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor(11, 37, 69)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333

    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 18, 10),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = doc.styles[name]
        style.font.name = FONT_NAME
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def build():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    configure_styles(doc)

    header = section.header.paragraphs[0]
    header.text = "全球教育公平数据可视化分析｜课堂展示逐字稿"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph_font(header, size=9, color="64748B")

    footer = section.footer.paragraphs[0]
    footer.text = "演讲时长建议：8-10 分钟"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_font(footer, size=9, color="64748B")

    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_after = Pt(3)
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run(TITLE)
    title_run.font.name = FONT_NAME
    title_run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    title_run.font.size = Pt(26)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(31, 95, 191)

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_p.paragraph_format.space_after = Pt(18)
    sub_run = sub_p.add_run(SUBTITLE)
    sub_run.font.name = FONT_NAME
    sub_run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    sub_run.font.size = Pt(13)
    sub_run.font.color.rgb = RGBColor(67, 83, 107)

    table = doc.add_table(rows=3, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    widths = [Inches(1.4), Inches(4.9)]
    entries = [
        ("展示主题", "构建更公平的教育未来：全球教育公平数据可视化分析"),
        ("展示结构", "投入与师资 → 机会与平等 → 质量与成果"),
        ("使用方式", "正文可直接照读；浅色提示行为网页操作提示，不需要读出口。"),
    ]
    for row, (label, value) in zip(table.rows, entries):
        for idx, cell in enumerate(row.cells):
            cell.width = widths[idx]
            set_cell_shading(cell, "F4F8FC" if idx == 0 else "FFFFFF")
        set_cell_text(row.cells[0], label, bold=True)
        set_cell_text(row.cells[1], value)

    add_para(
        doc,
        "提示：正式展示时，建议先用首页建立主题，再顺着左侧目录依次进入三个模块。讲图表时先说明“看什么”，再解释“发现什么”，最后落到“说明什么问题”。",
        size=10.5,
        color="43536B",
        italic=True,
        after=12,
    )

    for title, paragraphs, instructions in SECTIONS:
        doc.add_heading(title, level=1)
        for item in instructions:
            add_instruction(doc, item)
        for para in paragraphs:
            add_para(doc, para)

    doc.save(OUT)


if __name__ == "__main__":
    build()
